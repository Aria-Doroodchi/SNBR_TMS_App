"""
Parse motor nerve-conduction study files (.pdf / .docx) and extract CMAP data.

Each source file describes one participant-visit and contains a table of
nerve-site / muscle / latency / amplitude rows. Participant ID is extracted
from the filename (per lab convention); Visit Date is taken from the first
line of the document body.

Returns lists of plain dicts (no DataFrames). Downstream modules such as
``processing.df_builder`` are responsible for converting these records into
DataFrames and merging them with the main MEM DataFrame.

Public API
----------
cmap_output_columns()          -> list[str]
parse_cmap_file(filepath)      -> dict
parse_cmap_directory(input_dir) -> list[dict]
"""

from __future__ import annotations

import json
import re
import warnings
from datetime import datetime
from pathlib import Path

from parser.mem_parser import iter_files, normalize_dirs

_STUDY_ID_PATTERN = re.compile(r"([A-Za-z]+)\d*-0*(\d+)", flags=re.IGNORECASE)
_VISIT_DATE_PATTERN = re.compile(r"Visit Date:\s*([0-9A-Za-z/:\-\s]+?)\s*$", re.IGNORECASE)

# Date formats seen in source files. First match wins.
_DATE_FORMATS = (
    "%d-%b-%y %I:%M %p",   # EMGRQ PDF:  "14-Apr-26 9:26 AM"
    "%d-%b-%Y %I:%M %p",   # EMGRQ PDF (4-digit year)
    "%m/%d/%Y %I:%M %p",   # Natus DOCX: "3/27/2025 12:27 PM"
    "%m/%d/%y %I:%M %p",
    "%d/%m/%Y %I:%M %p",
    "%d-%b-%y",
    "%d-%b-%Y",
    "%m/%d/%Y",
    "%d/%m/%Y",
)


def cmap_output_columns() -> list[str]:
    """Return the stable output schema for parsed CMAP records."""
    return ["Study", "ID", "Date", "CMAP_table", "MUNIX_table", "source_file"]


# Logical column keys for the MUNIX table and the header tokens we accept for
# each. Matching is case-insensitive and ignores whitespace and the leading
# ``#``/``%`` glyphs.
_MUNIX_COLUMN_ALIASES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("num_sip",  ("sip", "nsip", "numsip", "#sip")),
    ("a",        ("a",)),
    ("alpha",    ("alpha",)),
    ("munix",    ("munix",)),
    ("musix",    ("musix",)),
)


# ---------------------------------------------------------------------------
# Shared extraction helpers
# ---------------------------------------------------------------------------

def _extract_study_and_id(text: str | None) -> tuple[str | None, int | None]:
    """Extract (study_name, participant_id) using the same regex MEM/CSP use."""
    if text is None:
        return None, None
    match = _STUDY_ID_PATTERN.search(str(text))
    if match:
        return match.group(1).upper(), int(match.group(2))
    return None, None


def _parse_visit_date(raw: str | None) -> str | None:
    """Parse a Visit Date string and normalise it to ``dd/mm/YYYY``."""
    if not raw:
        return None
    candidate = raw.strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(candidate, fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return None


def _extract_visit_date_from_text(text: str) -> str | None:
    """Scan *text* for a ``Visit Date: ...`` line and normalise the value."""
    for line in text.splitlines():
        line = line.strip()
        if line.lower().startswith("visit date:"):
            raw = line.split(":", 1)[1].strip()
            # Handle tab-separated docx paragraphs like "Visit Date:\t3/27/2025 12:27 PM"
            raw = raw.replace("\t", " ").strip()
            parsed = _parse_visit_date(raw)
            if parsed:
                return parsed
    return None


def _to_float(value) -> float | None:
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    try:
        return float(s)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Table row extraction (shared by both formats)
# ---------------------------------------------------------------------------

def _locate_columns(header: list[str]) -> dict[str, int] | None:
    """Map logical column names → positional index from a header row."""
    lookup: dict[str, int] = {}
    for idx, raw in enumerate(header):
        if raw is None:
            continue
        key = str(raw).strip().lower().split("\n")[0]
        if key.startswith("nerve"):
            lookup["nerve_site"] = idx
        elif key.startswith("muscle"):
            lookup["muscle"] = idx
        elif key.startswith("latency"):
            lookup["latency"] = idx
        elif key.startswith("amplitude"):
            lookup["amplitude"] = idx
    required = {"nerve_site", "muscle", "latency", "amplitude"}
    return lookup if required.issubset(lookup) else None


def _is_group_row(cells: list[str | None], nerve_idx: int) -> bool:
    """Group rows have the nerve text in the first cell and the rest blank/equal."""
    first = (cells[nerve_idx] or "").strip()
    if not first:
        return False
    others = [(c or "").strip() for i, c in enumerate(cells) if i != nerve_idx]
    # PDF: other cells empty. DOCX: other cells equal to first (merged fill).
    if all(v == "" for v in others):
        return True
    return all(v == first for v in others if v != "")


def _rows_from_table(
    table: list[list[str | None]],
) -> list[dict]:
    """Walk a 2-D table and return CMAP rows preserving source order."""
    if not table:
        return []

    header = [("" if c is None else str(c)) for c in table[0]]
    cols = _locate_columns(header)
    if cols is None:
        return []

    nerve_idx = cols["nerve_site"]
    muscle_idx = cols["muscle"]
    latency_idx = cols["latency"]
    amp_idx = cols["amplitude"]

    rows: list[dict] = []
    current_group = ""
    for raw in table[1:]:
        cells = [("" if c is None else str(c).strip()) for c in raw]
        # Pad short rows so indexing doesn't blow up
        while len(cells) <= max(nerve_idx, muscle_idx, latency_idx, amp_idx):
            cells.append("")

        # Units row (e.g., ["", "", "ms", "mV", ...]) — skip if the nerve cell
        # is empty and every populated cell is a unit token.
        if not cells[nerve_idx] and all(
            c == "" or c.lower() in {"ms", "mv", "%", "°c", "cm", "mv/d", "ms/d"}
            for c in cells
        ):
            continue

        if _is_group_row(cells, nerve_idx):
            current_group = cells[nerve_idx]
            continue

        latency = _to_float(cells[latency_idx])
        amplitude = _to_float(cells[amp_idx])
        if latency is None and amplitude is None:
            continue

        site = cells[nerve_idx]
        nerve_site = f"{current_group} / {site}" if current_group and site else (current_group or site)
        rows.append({
            "nerve_site": nerve_site,
            "muscle": cells[muscle_idx],
            "latency_ms": latency,
            "amplitude_mv": amplitude,
        })
    return rows


# ---------------------------------------------------------------------------
# MUNIX table extraction
# ---------------------------------------------------------------------------

def _normalize_header_token(value) -> str:
    """Lowercase, strip whitespace, ``#`` and ``%`` from a header cell."""
    if value is None:
        return ""
    s = str(value).strip().lower().split("\n")[0]
    return s.replace("#", "").replace("%", "").replace(" ", "")


def _locate_munix_columns(header: list[str]) -> dict[str, int] | None:
    """Map MUNIX logical columns → positional index from a header row.

    Returns ``None`` unless **all five** MUNIX columns are present.
    """
    lookup: dict[str, int] = {}
    for idx, raw in enumerate(header):
        token = _normalize_header_token(raw)
        if not token:
            continue
        for key, aliases in _MUNIX_COLUMN_ALIASES:
            if token in aliases:
                lookup.setdefault(key, idx)
                break
    required = {k for k, _ in _MUNIX_COLUMN_ALIASES}
    return lookup if required.issubset(lookup) else None


def _munix_rows_from_table(table: list[list[str | None]]) -> list[dict]:
    """Extract MUNIX data rows from a 2-D table, or [] if the header doesn't match."""
    if not table:
        return []
    header = [("" if c is None else str(c)) for c in table[0]]
    cols = _locate_munix_columns(header)
    if cols is None:
        return []

    out: list[dict] = []
    for raw in table[1:]:
        cells = [("" if c is None else str(c).strip()) for c in raw]
        if not cells:
            continue
        # Pad short rows so indexing never goes out of bounds.
        max_idx = max(cols.values())
        while len(cells) <= max_idx:
            cells.append("")
        # Skip rows where every MUNIX column is blank
        if not any(cells[cols[k]] for k, _ in _MUNIX_COLUMN_ALIASES):
            continue
        row: dict = {}
        for key, _aliases in _MUNIX_COLUMN_ALIASES:
            raw_val = cells[cols[key]]
            num = _to_float(raw_val)
            row[key] = num if num is not None else (raw_val or None)
        out.append(row)
    return out


# ---------------------------------------------------------------------------
# Format-specific readers
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> tuple[str | None, list[dict], list[dict]]:
    """Extract ``(visit_date, cmap_rows, munix_rows)`` from an EMGRQ-style PDF."""
    import pdfplumber

    visit_date: str | None = None
    cmap: list[dict] = []
    munix: list[dict] = []
    with pdfplumber.open(path) as pdf:
        if not pdf.pages:
            return None, [], []
        first = pdf.pages[0]
        text = first.extract_text() or ""
        visit_date = _extract_visit_date_from_text(text)
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                m = _munix_rows_from_table(table)
                if m:
                    munix.extend(m)
                    continue
                c = _rows_from_table(table)
                if c:
                    cmap.extend(c)
    return visit_date, cmap, munix


def _parse_docx(path: Path) -> tuple[str | None, list[dict], list[dict]]:
    """Extract ``(visit_date, cmap_rows, munix_rows)`` from a Natus-style DOCX."""
    from docx import Document

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    visit_date = _extract_visit_date_from_text(text)

    cmap: list[dict] = []
    munix: list[dict] = []
    for tbl in doc.tables:
        grid: list[list[str]] = []
        for row in tbl.rows:
            grid.append([cell.text.strip() for cell in row.cells])
        m = _munix_rows_from_table(grid)
        if m:
            munix.extend(m)
            continue
        c = _rows_from_table(grid)
        if c:
            cmap.extend(c)
    return visit_date, cmap, munix


# ---------------------------------------------------------------------------
# Public parser entry points
# ---------------------------------------------------------------------------

def parse_cmap_file(filepath: str | Path) -> dict:
    """Parse a single CMAP file and return a record dict.

    The returned dict uses the schema from :func:`cmap_output_columns`.
    ``CMAP_table`` holds a JSON string (list of row dicts) so the value can
    round-trip through CSV. Missing fields are ``None``.
    """
    path = Path(filepath)
    suffix = path.suffix.lower()

    study, pid = _extract_study_and_id(path.stem)

    try:
        if suffix == ".pdf":
            visit_date, cmap_rows, munix_rows = _parse_pdf(path)
        elif suffix == ".docx":
            visit_date, cmap_rows, munix_rows = _parse_docx(path)
        else:
            warnings.warn(f"Unsupported CMAP file extension: {path.name}")
            visit_date, cmap_rows, munix_rows = None, [], []
    except Exception as exc:  # pragma: no cover — defensive
        warnings.warn(f"Failed to parse CMAP file {path.name}: {exc}")
        visit_date, cmap_rows, munix_rows = None, [], []

    return {
        "Study": study,
        "ID": pid,
        "Date": visit_date,
        "CMAP_table": json.dumps(cmap_rows) if cmap_rows else None,
        "MUNIX_table": json.dumps(munix_rows) if munix_rows else None,
        "source_file": path.name,
    }


def parse_cmap_directory(
    input_dir: str | Path | list[str | Path] | None,
) -> list[dict]:
    """Parse every .pdf/.docx file in *input_dir* (recursive) and return records.

    *input_dir* may be a single directory or a list of directories.  Files
    that yield no usable rows are silently skipped. Temporary Word lock
    files (``~$...docx``) are ignored.
    """
    if not normalize_dirs(input_dir):
        return []

    records: list[dict] = []
    for path in iter_files(input_dir, "*"):
        if path.name.startswith("~$"):
            continue
        if path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        record = parse_cmap_file(path)
        if record.get("ID") is None:
            continue
        if not record.get("CMAP_table") and not record.get("MUNIX_table"):
            continue
        records.append(record)
    return records
